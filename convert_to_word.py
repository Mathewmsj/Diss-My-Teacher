#!/usr/bin/env python3
"""
Convert Markdown essay to Word document
"""
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_code_block(paragraph, code_text):
    """Add code block with monospace font"""
    run = paragraph.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

def parse_markdown_to_word(md_file, docx_file):
    """Convert Markdown file to Word document"""
    doc = Document()
    
    # Set default font to Times New Roman
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set margins (IEEE format: 1 inch margins)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                if code_block_lines:
                    p = doc.add_paragraph()
                    p.style = 'No Spacing'
                    add_code_block(p, '\n'.join(code_block_lines))
                    doc.add_paragraph()  # Add spacing after code
                code_block_lines = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Handle headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        
        # Handle horizontal rules
        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.add_run('_' * 50).italic = True
        
        # Handle tables
        elif '|' in line and line.strip().startswith('|'):
            table_data = []
            while i < len(lines) and '|' in lines[i]:
                row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                if row and not all(c == '-' for c in ''.join(row)):
                    table_data.append(row)
                i += 1
            i -= 1
            
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Light Grid Accent 1'
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_data in enumerate(row_data):
                        table.rows[row_idx].cells[col_idx].text = cell_data
        
        # Handle lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line.strip()):
            p = doc.add_paragraph(line.strip(), style='List Number')
        
        # Handle bold text
        elif '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        
        # Handle regular paragraphs
        elif line.strip():
            # Check for inline code
            if '`' in line:
                p = doc.add_paragraph()
                parts = re.split(r'(`[^`]+`)', line)
                for part in parts:
                    if part.startswith('`') and part.endswith('`'):
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                    else:
                        p.add_run(part)
            else:
                p = doc.add_paragraph(line.strip())
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"Word document saved to: {docx_file}")

if __name__ == '__main__':
    parse_markdown_to_word('Login_Implementation_Essay.md', 'Login_Implementation_Essay.docx')
