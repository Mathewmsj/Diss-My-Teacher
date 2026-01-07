#!/usr/bin/env python3
"""
生成Word文档的脚本
需要安装: pip install python-docx
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("需要安装 python-docx: pip install python-docx")
    exit(1)

# 创建文档
doc = Document()

# 设置中文字体
def set_chinese_font(run):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 标题
title = doc.add_heading('Deployment of Rate My Teacher Application: A Step-by-Step Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 作者信息
author_para = doc.add_paragraph()
author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_run = author_para.add_run('Author: Mathew\nDate: January 2026')
set_chinese_font(author_run)
author_run.font.size = Pt(12)

doc.add_paragraph()  # 空行

# Abstract
doc.add_heading('Abstract', 1)
abstract_para = doc.add_paragraph()
abstract_run = abstract_para.add_run(
    'This document presents a comprehensive guide for deploying the Rate My Teacher web application on a Linux server. '
    'The deployment process covers both IP-based access and domain-based access configurations. '
    'The application consists of a Django REST API backend and a Vue.js frontend, deployed using a unified startup script that manages both services simultaneously.'
)
set_chinese_font(abstract_run)
abstract_run.italic = True

doc.add_page_break()

# I. INTRODUCTION
doc.add_heading('I. INTRODUCTION', 1)
intro_para = doc.add_paragraph()
intro_run = intro_para.add_run(
    'The Rate My Teacher application is a full-stack web application designed to allow students to rate and review teachers. '
    'The system architecture comprises:'
)
set_chinese_font(intro_run)

# 列表
items = [
    'Backend: Django 4.2.7 REST API with SQLite database',
    'Frontend: Vue.js 3 with Vite build tool',
    'Deployment Platform: Linux server with nginx reverse proxy'
]
for item in items:
    para = doc.add_paragraph(item, style='List Bullet')
    for run in para.runs:
        set_chinese_font(run)

para = doc.add_paragraph()
run = para.add_run(
    'This paper documents the complete deployment process, including server setup, code deployment, service configuration, '
    'and access methods through both IP addresses and custom domains.'
)
set_chinese_font(run)

doc.add_page_break()

# II. DEPLOYMENT METHODOLOGY
doc.add_heading('II. DEPLOYMENT METHODOLOGY', 1)

doc.add_heading('A. Prerequisites', 2)
para = doc.add_paragraph('Before deployment, ensure the following prerequisites are met:')
for run in para.runs:
    set_chinese_font(run)

prereq_items = [
    'Server Access: SSH access to the Linux server (IP: 110.40.153.38)',
    'Python Environment: Python 3.9 or higher installed',
    'Node.js: Node.js v22+ installed for frontend dependencies',
    'Git: Git installed for code version control',
    'Domain Configuration: Domain mathew.yunguhs.com configured in nginx'
]
for item in prereq_items:
    para = doc.add_paragraph(item, style='List Number')
    for run in para.runs:
        set_chinese_font(run)

doc.add_heading('B. Step-by-Step Deployment Process', 2)

# Step 1
doc.add_heading('Step 1: Server Connection', 3)
para = doc.add_paragraph('Connect to the server using SSH with your username:')
for run in para.runs:
    set_chinese_font(run)
code_para = doc.add_paragraph('ssh mathew@110.40.153.38', style='No Spacing')
code_para.style.font.name = 'Courier New'
code_para.paragraph_format.left_indent = Inches(0.5)

# Step 2
doc.add_heading('Step 2: Clone Repository', 3)
para = doc.add_paragraph('Clone the project repository from GitHub:')
for run in para.runs:
    set_chinese_font(run)
code_para = doc.add_paragraph('git clone https://github.com/Mathewmsj/Diss-My-Teacher.git\ncd Diss-My-Teacher', style='No Spacing')
code_para.style.font.name = 'Courier New'
code_para.paragraph_format.left_indent = Inches(0.5)

para = doc.add_paragraph('If the repository already exists, update it:')
for run in para.runs:
    set_chinese_font(run)
code_para = doc.add_paragraph('cd Diss-My-Teacher\ngit pull', style='No Spacing')
code_para.style.font.name = 'Courier New'
code_para.paragraph_format.left_indent = Inches(0.5)

# Step 3
doc.add_heading('Step 3: Backend Setup', 3)
para = doc.add_paragraph('Navigate to the backend directory and create a Python virtual environment:')
for run in para.runs:
    set_chinese_font(run)
code_para = doc.add_paragraph('cd backend\npython3.9 -m venv backend-env\nsource backend-env/bin/activate', style='No Spacing')
code_para.style.font.name = 'Courier New'
code_para.paragraph_format.left_indent = Inches(0.5)

para = doc.add_paragraph('Install Python dependencies:')
for run in para.runs:
    set_chinese_font(run)
code_para = doc.add_paragraph('pip install --upgrade pip setuptools wheel\npip install -r requirements.txt', style='No Spacing')
code_para.style.font.name = 'Courier New'
code_para.paragraph_format.left_indent = Inches(0.5)

para = doc.add_paragraph('The requirements include:')
for run in para.runs:
    set_chinese_font(run)
req_items = [
    'Django==4.2.7',
    'djangorestframework==3.14.0',
    'django-cors-headers==4.6.0',
    'gunicorn==21.2.0',
    'dj-database-url==2.3.0',
    'psycopg2-binary==2.9.10'
]
for item in req_items:
    para = doc.add_paragraph(item, style='List Bullet')
    for run in para.runs:
        set_chinese_font(run)

# Step 4
doc.add_heading('Step 4: Database Migration', 3)
para = doc.add_paragraph('Run database migrations to set up the database schema:')
for run in para.runs:
    set_chinese_font(run)
code_para = doc.add_paragraph('python manage.py migrate', style='No Spacing')
code_para.style.font.name = 'Courier New'
code_para.paragraph_format.left_indent = Inches(0.5)

para = doc.add_paragraph('Create a superuser account for administrative access:')
for run in para.runs:
    set_chinese_font(run)
code_para = doc.add_paragraph('python create_superadmin.py', style='No Spacing')
code_para.style.font.name = 'Courier New'
code_para.paragraph_format.left_indent = Inches(0.5)

para = doc.add_paragraph('This creates a superadmin account with:')
for run in para.runs:
    set_chinese_font(run)
admin_items = [
    'Username: Starry',
    'Password: msj20070528'
]
for item in admin_items:
    para = doc.add_paragraph(item, style='List Bullet')
    for run in para.runs:
        set_chinese_font(run)

# Step 5
doc.add_heading('Step 5: Frontend Setup', 3)
para = doc.add_paragraph('Return to the project root directory and install Node.js dependencies:')
for run in para.runs:
    set_chinese_font(run)
code_para = doc.add_paragraph('cd ..\nnpm install', style='No Spacing')
code_para.style.font.name = 'Courier New'
code_para.paragraph_format.left_indent = Inches(0.5)

# Step 6
doc.add_heading('Step 6: Configuration Verification', 3)
para = doc.add_paragraph('Verify the following configuration files:')
for run in para.runs:
    set_chinese_font(run)

para = doc.add_paragraph('Backend Configuration (backend/backend/settings.py):')
for run in para.runs:
    set_chinese_font(run)
    run.bold = True
config_items = [
    'ALLOWED_HOSTS = [\'*\'] - Allows access from any host',
    'Database configured for SQLite by default',
    'CORS headers enabled for cross-origin requests'
]
for item in config_items:
    para = doc.add_paragraph(item, style='List Bullet')
    for run in para.runs:
        set_chinese_font(run)

para = doc.add_paragraph('Frontend Configuration (vite.config.js):')
for run in para.runs:
    set_chinese_font(run)
    run.bold = True
frontend_config = [
    'Server host set to 0.0.0.0 for external access',
    'Port configuration via environment variable PORT'
]
for item in frontend_config:
    para = doc.add_paragraph(item, style='List Bullet')
    for run in para.runs:
        set_chinese_font(run)

para = doc.add_paragraph('API Configuration (src/api.js):')
for run in para.runs:
    set_chinese_font(run)
    run.bold = True
api_config = [
    'Automatic detection of access method (IP vs domain)',
    'IP access: Uses http://IP:PORT/api',
    'Domain access: Uses relative path /api (handled by nginx)'
]
for item in api_config:
    para = doc.add_paragraph(item, style='List Bullet')
    for run in para.runs:
        set_chinese_font(run)

doc.add_page_break()

# III. SERVICE DEPLOYMENT
doc.add_heading('III. SERVICE DEPLOYMENT', 1)

doc.add_heading('A. IP-Based Access Deployment', 2)
para = doc.add_paragraph('For IP-based access, the application uses default ports:')
for run in para.runs:
    set_chinese_font(run)
port_items = [
    'Backend: Port 5009',
    'Frontend: Port 5010'
]
for item in port_items:
    para = doc.add_paragraph(item, style='List Bullet')
    for run in para.runs:
        set_chinese_font(run)

doc.add_heading('Starting Services', 3)
para = doc.add_paragraph('Execute the startup script:')
for run in para.runs:
    set_chinese_font(run)
code_para = doc.add_paragraph('./start.sh', style='No Spacing')
code_para.style.font.name = 'Courier New'
code_para.paragraph_format.left_indent = Inches(0.5)

para = doc.add_paragraph('Or specify custom ports:')
for run in para.runs:
    set_chinese_font(run)
code_para = doc.add_paragraph('./start.sh 5009 5010', style='No Spacing')
code_para.style.font.name = 'Courier New'
code_para.paragraph_format.left_indent = Inches(0.5)

para = doc.add_paragraph('The startup script performs the following operations:')
for run in para.runs:
    set_chinese_font(run)
script_ops = [
    'Port Cleanup: Checks and terminates any processes using the target ports',
    'Backend Launch: Starts Django development server on 0.0.0.0:5009',
    'Frontend Launch: Starts Vite development server on 0.0.0.0:5010',
    'Process Management: Saves process IDs to backend.pid and frontend.pid'
]
for item in script_ops:
    para = doc.add_paragraph(item, style='List Number')
    for run in para.runs:
        set_chinese_font(run)

doc.add_heading('Accessing the Application', 3)
para = doc.add_paragraph('After successful startup, access the application via:')
for run in para.runs:
    set_chinese_font(run)
access_items = [
    'Frontend: http://110.40.153.38:5010',
    'Backend API: http://110.40.153.38:5009/api'
]
for item in access_items:
    para = doc.add_paragraph(item, style='List Bullet')
    for run in para.runs:
        set_chinese_font(run)

doc.add_heading('B. Domain-Based Access Deployment', 2)
para = doc.add_paragraph('For domain-based access, the application uses assigned ports:')
for run in para.runs:
    set_chinese_font(run)
domain_port_items = [
    'Backend: Port 8806',
    'Frontend: Port 8807',
    'Domain: mathew.yunguhs.com'
]
for item in domain_port_items:
    para = doc.add_paragraph(item, style='List Bullet')
    for run in para.runs:
        set_chinese_font(run)

doc.add_heading('Starting Services with Domain Ports', 3)
para = doc.add_paragraph('Execute the startup script with domain-specific ports:')
for run in para.runs:
    set_chinese_font(run)
code_para = doc.add_paragraph('./start.sh 8806 8807', style='No Spacing')
code_para.style.font.name = 'Courier New'
code_para.paragraph_format.left_indent = Inches(0.5)

doc.add_heading('Accessing the Application', 3)
para = doc.add_paragraph('After successful startup, access the application via:')
for run in para.runs:
    set_chinese_font(run)
domain_access_items = [
    'Frontend: http://mathew.yunguhs.com or https://mathew.yunguhs.com',
    'Backend API: http://mathew.yunguhs.com/api or https://mathew.yunguhs.com/api'
]
for item in domain_access_items:
    para = doc.add_paragraph(item, style='List Bullet')
    for run in para.runs:
        set_chinese_font(run)

doc.add_page_break()

# IV. OPERATIONAL PROCEDURES
doc.add_heading('IV. OPERATIONAL PROCEDURES', 1)

doc.add_heading('A. Service Management', 2)

doc.add_heading('Stopping Services', 3)
para = doc.add_paragraph('Stop all services using:')
for run in para.runs:
    set_chinese_font(run)
code_para = doc.add_paragraph('./stop.sh', style='No Spacing')
code_para.style.font.name = 'Courier New'
code_para.paragraph_format.left_indent = Inches(0.5)

doc.add_heading('Restarting Services', 3)
para = doc.add_paragraph('Restart services using:')
for run in para.runs:
    set_chinese_font(run)
code_para = doc.add_paragraph('./restart.sh', style='No Spacing')
code_para.style.font.name = 'Courier New'
code_para.paragraph_format.left_indent = Inches(0.5)

doc.add_heading('Viewing Logs', 3)
para = doc.add_paragraph('Monitor service logs in real-time:')
for run in para.runs:
    set_chinese_font(run)
code_para = doc.add_paragraph('# Backend logs\ntail -f backend.log\n\n# Frontend logs\ntail -f frontend.log', style='No Spacing')
code_para.style.font.name = 'Courier New'
code_para.paragraph_format.left_indent = Inches(0.5)

doc.add_page_break()

# V. SECURITY CONSIDERATIONS
doc.add_heading('V. SECURITY CONSIDERATIONS', 1)

doc.add_heading('A. Network Security', 2)
security_items = [
    'Firewall Configuration: Ensure ports 5009, 5010, 8806, and 8807 are accessible',
    'HTTPS: Use HTTPS for domain access when SSL certificates are configured',
    'CORS: Configured to allow cross-origin requests from frontend domain'
]
for item in security_items:
    para = doc.add_paragraph(item, style='List Number')
    for run in para.runs:
        set_chinese_font(run)

doc.add_heading('B. Application Security', 2)
app_security_items = [
    'Secret Key: Change SECRET_KEY in production settings',
    'Debug Mode: Set DEBUG = False in production',
    'Allowed Hosts: Configure specific hosts instead of [\'*\'] in production',
    'Authentication: Token-based authentication for API access'
]
for item in app_security_items:
    para = doc.add_paragraph(item, style='List Number')
    for run in para.runs:
        set_chinese_font(run)

doc.add_page_break()

# VI. RESULTS AND VALIDATION
doc.add_heading('VI. RESULTS AND VALIDATION', 1)

doc.add_heading('A. IP-Based Access Validation', 2)
para = doc.add_paragraph('After deployment, validate IP-based access:')
for run in para.runs:
    set_chinese_font(run)

validation_items = [
    'Frontend Access: Navigate to http://110.40.153.38:5010 - Expected: Vue.js application loads successfully',
    'Backend API Access: Test API endpoint http://110.40.153.38:5009/api/healthz - Expected: Returns {"status": "ok"}',
    'Service Status: Run ./check.sh - Expected: Both services running, ports listening on 0.0.0.0'
]
for item in validation_items:
    para = doc.add_paragraph(item, style='List Number')
    for run in para.runs:
        set_chinese_font(run)

doc.add_heading('B. Domain-Based Access Validation', 2)
para = doc.add_paragraph('After deployment, validate domain-based access:')
for run in para.runs:
    set_chinese_font(run)

domain_validation_items = [
    'Frontend Access: Navigate to https://mathew.yunguhs.com - Expected: Vue.js application loads successfully',
    'Backend API Access: Test API endpoint https://mathew.yunguhs.com/api/healthz - Expected: Returns {"status": "ok"}',
    'Root Path Redirect: Navigate to https://mathew.yunguhs.com/api/ - Expected: Redirects to https://mathew.yunguhs.com'
]
for item in domain_validation_items:
    para = doc.add_paragraph(item, style='List Number')
    for run in para.runs:
        set_chinese_font(run)

doc.add_page_break()

# VII. CONCLUSION
doc.add_heading('VII. CONCLUSION', 1)
para = doc.add_paragraph(
    'This deployment guide provides a comprehensive methodology for deploying the Rate My Teacher application on a Linux server. '
    'The deployment supports both IP-based and domain-based access methods, with automatic configuration detection and routing.'
)
for run in para.runs:
    set_chinese_font(run)

para = doc.add_paragraph('Key achievements:')
for run in para.runs:
    set_chinese_font(run)
    run.bold = True

achievements = [
    'Unified Deployment: Single script manages both backend and frontend services',
    'Flexible Access: Supports both IP and domain access methods',
    'Automated Configuration: Frontend automatically detects access method and configures API endpoints',
    'Service Management: Comprehensive scripts for start, stop, restart, and monitoring'
]
for item in achievements:
    para = doc.add_paragraph(item, style='List Number')
    for run in para.runs:
        set_chinese_font(run)

para = doc.add_paragraph(
    'The deployment process is reproducible and can be executed in approximately 10-15 minutes, '
    'including dependency installation and database migration.'
)
for run in para.runs:
    set_chinese_font(run)

doc.add_page_break()

# REFERENCES
doc.add_heading('REFERENCES', 1)
references = [
    'Django Software Foundation, "Django Documentation," https://docs.djangoproject.com/, accessed Jan. 2026.',
    'Vue.js Team, "Vite Documentation," https://vitejs.dev/, accessed Jan. 2026.',
    'Nginx Inc., "Nginx Documentation," https://nginx.org/en/docs/, accessed Jan. 2026.',
    'Python Software Foundation, "Python Virtual Environments," https://docs.python.org/3/tutorial/venv.html, accessed Jan. 2026.',
    'Node.js Foundation, "Node.js Documentation," https://nodejs.org/docs/, accessed Jan. 2026.'
]
for ref in references:
    para = doc.add_paragraph(ref, style='List Number')
    for run in para.runs:
        set_chinese_font(run)

doc.add_page_break()

# APPENDICES
doc.add_heading('Appendix A: File Structure', 2)
code_para = doc.add_paragraph(
    'Diss-My-Teacher/\n'
    '├── backend/\n'
    '│   ├── backend/\n'
    '│   │   ├── settings.py      # Django settings\n'
    '│   │   └── urls.py          # URL routing\n'
    '│   ├── api/                 # API application\n'
    '│   ├── manage.py\n'
    '│   ├── requirements.txt\n'
    '│   ├── backend-env/         # Python virtual environment\n'
    '│   └── db.sqlite3           # SQLite database\n'
    '├── src/                     # Frontend source code\n'
    '├── start.sh                 # Startup script\n'
    '├── stop.sh                  # Stop script\n'
    '├── check.sh                 # Diagnostic script\n'
    '└── package.json             # Node.js dependencies',
    style='No Spacing'
)
code_para.style.font.name = 'Courier New'
code_para.paragraph_format.left_indent = Inches(0.5)

doc.add_heading('Appendix B: Port Configuration', 2)
table = doc.add_table(rows=3, cols=3)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = 'Service'
header_cells[1].text = 'IP Access'
header_cells[2].text = 'Domain Access'
for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_chinese_font(run)
            run.bold = True

data_rows = [
    ['Backend', '5009', '8806'],
    ['Frontend', '5010', '8807']
]
for i, row_data in enumerate(data_rows, 1):
    row_cells = table.rows[i].cells
    for j, cell_text in enumerate(row_data):
        row_cells[j].text = cell_text
        for paragraph in row_cells[j].paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run)

doc.add_heading('Appendix C: Environment Variables', 2)
env_vars = [
    'DATABASE_URL: PostgreSQL connection string (optional, for Render database)',
    'FRONTEND_URL: Frontend URL for redirects (optional)',
    'PORT: Frontend port (default: 8080, overridden by script)'
]
for var in env_vars:
    para = doc.add_paragraph(var, style='List Bullet')
    for run in para.runs:
        set_chinese_font(run)

# 保存文档
output_file = 'DEPLOYMENT_ESSAY.docx'
doc.save(output_file)
print(f'Word文档已生成: {output_file}')

